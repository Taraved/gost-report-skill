#!/usr/bin/env python3
"""GOST 7.32-2001 Lab Report .docx Generator.

Reads structured JSON data and produces a formatted Word document
compliant with the GOST 7.32-2001 standard for academic reports.

Typical usage:
    python3 scripts/generate_docx.py --input data.json --output report.docx

The input JSON schema is documented in the project README and skill/SKILL.md.
"""

from __future__ import annotations

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Any, Optional

from docx import Document  # type: ignore[import-untyped]
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]
from docx.oxml import OxmlElement  # type: ignore[import-untyped]
from docx.oxml.ns import qn  # type: ignore[import-untyped]
from docx.shared import Cm, Pt, RGBColor  # type: ignore[import-untyped]

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

DEFAULT_HEADER_LINES: tuple[str, ...] = (
    "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ",
    "ФГАОУ ВО Уральский федеральный университет",
    "Институт радиоэлектроники и информационных технологий-РТФ",
)

FONT_NAME: str = "Times New Roman"
FONT_SIZE: Pt = Pt(14)
CODE_FONT: str = "Roboto Mono"
CODE_FONT_SIZE: Pt = Pt(11)
LINE_SPACING: float = 1.5
INDENT_CM: float = 1.25

MARGIN_LEFT_CM: float = 3.0
MARGIN_RIGHT_CM: float = 1.5
MARGIN_TOP_CM: float = 2.0
MARGIN_BOTTOM_CM: float = 2.0

REQUIRED_JSON_KEYS: tuple[str, ...] = ("lab_number", "tasks")


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------


def _set_run_font(
    run: Any,
    font_name: str = FONT_NAME,
    font_size: Pt = FONT_SIZE,
    bold: bool = False,
) -> None:
    """Apply font properties to a ``run`` element, including Cyrillic-encoding fields.

    Args:
        run: python-docx Run object.
        font_name: Font family name (default: Times New Roman).
        font_size: Font size in points (default: 14pt).
        bold: Whether the text should be bold (default: False).
    """
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), font_name)
    rPr.append(rFonts)


def _fix_heading_fonts(
    style_name: str,
    alignment: Optional[int] = None,
    indent: bool = True,
) -> None:
    """Override font and paragraph settings for a built-in heading style.

    Args:
        style_name: Name of the style in the document (e.g. 'Heading 1').
        alignment: Paragraph alignment constant from ``WD_ALIGN_PARAGRAPH``.
        indent: Whether to apply a first-line indent (default: True).
    """
    doc_styles: Any = None  # set by caller
    # NOTE: this is injected via closure; refactored to top-level for clarity.
    # We pass ``doc`` explicitly instead.  See ``setup_gost_styles``.


def _add_styled_paragraph(
    doc: Document,
    text: str,
    style: Optional[str] = None,
    alignment: Optional[int] = None,
    bold: bool = False,
    space_before: float = 0,
    space_after: float = 0,
    no_indent: bool = False,
) -> Any:
    """Add a paragraph with consistent font and spacing settings.

    Args:
        doc: The active Document instance.
        text: Paragraph text content.
        style: Name of a Word style to apply (e.g. 'Heading 1').
        alignment: ``WD_ALIGN_PARAGRAPH`` constant.
        bold: Render text in bold.
        space_before: Additional space before the paragraph (in points).
        space_after: Additional space after the paragraph (in points).
        no_indent: Remove the standard first-line indent.

    Returns:
        The created paragraph object.
    """
    p = doc.add_paragraph(text, style=style)
    if alignment is not None:
        p.alignment = alignment
    if space_before:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after:
        p.paragraph_format.space_after = Pt(space_after)
    if no_indent:
        p.paragraph_format.first_line_indent = 0
    for run in p.runs:
        _set_run_font(run, bold=bold)
    return p


def _add_code_paragraph(doc: Document, text: str) -> None:
    """Add a single line of monospaced source code.

    Args:
        doc: The active Document instance.
        text: A line of source code (including leading whitespace).
    """
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = 0
    run = p.add_run(text)
    _set_run_font(run, font_name=CODE_FONT, font_size=CODE_FONT_SIZE)


def _add_page_number(run: Any) -> None:
    """Insert a dynamic page-number field (Word field code PAGE).

    Args:
        run: The run element where the field code will be injected.
    """
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_char_begin)
    run._element.append(instr_text)
    run._element.append(fld_char_end)


# ---------------------------------------------------------------------------
# Document configuration
# ---------------------------------------------------------------------------


def setup_gost_styles(doc: Document) -> None:
    """Configure built-in Word styles to match GOST 7.32-2001 requirements.

    Applies Times New Roman 14pt, 1.5 line spacing, justified alignment,
    and a 1.25 cm first-line indent to the Normal style.  Also adjusts
    Heading 1-3 for Cyrillic font rendering.

    Args:
        doc: The active Document instance.
    """
    # Normal text
    style_n = doc.styles["Normal"]
    font = style_n.font
    font.name = FONT_NAME
    font.size = FONT_SIZE
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style_n.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(INDENT_CM)

    # Heading 1-3
    def _fix_style(style_name: str, alignment: Optional[int] = None, indent: bool = True) -> None:
        """Apply Times New Roman and spacing to a heading style.

        Args:
            style_name: Name of the style to modify.
            alignment: Paragraph alignment constant.
            indent: Whether to apply a first-line indent of 1.25 cm.
        """
        s = doc.styles[style_name]
        rPr = s._element.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rFonts.set(qn(attr), FONT_NAME)
        rPr.append(rFonts)
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.font.size = FONT_SIZE
        s.font.bold = False
        s.paragraph_format.line_spacing = LINE_SPACING
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after = Pt(6)
        if alignment is not None:
            s.paragraph_format.alignment = alignment
        s.paragraph_format.first_line_indent = Cm(INDENT_CM) if indent else Cm(0)

    if "Heading 1" in doc.styles:
        _fix_style("Heading 1", alignment=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    for i in range(2, 4):
        name = f"Heading {i}"
        if name in doc.styles:
            _fix_style(name, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True)


def configure_section(doc: Document) -> None:
    """Set page margins and add a centred page-number footer.

    Margins follow GOST 7.32-2001: left 30 mm, right 15 mm, top/bottom 20 mm.

    Args:
        doc: The active Document instance.
    """
    s = doc.sections[0]
    s.left_margin = Cm(MARGIN_LEFT_CM)
    s.right_margin = Cm(MARGIN_RIGHT_CM)
    s.top_margin = Cm(MARGIN_TOP_CM)
    s.bottom_margin = Cm(MARGIN_BOTTOM_CM)
    para = s.footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    _set_run_font(run, font_size=Pt(12))
    _add_page_number(run)


def create_toc(doc: Document) -> None:
    """Insert a Word table of contents field (headings 1-3).

    The TOC is rendered when the user opens the document and presses Ctrl+A → F9.

    Args:
        doc: The active Document instance.
    """
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = 0
    run = p.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._element.append(fld_char)
    run._element.append(instr_text)
    run._element.append(fld_char2)
    run._element.append(fld_char3)


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------


def build_title_page(doc: Document, data: dict[str, Any]) -> None:
    """Render the title page with university, discipline, student, and teacher info.

    Args:
        doc: The active Document instance.
        data: The full report dictionary loaded from JSON.
    """
    for line in DEFAULT_HEADER_LINES:
        _add_styled_paragraph(doc, line, alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True)
    doc.add_paragraph()
    disc = data.get("discipline", "")
    if disc:
        _add_styled_paragraph(doc, f'«{disc}»', alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True)
    doc.add_paragraph()
    _add_styled_paragraph(doc, "ОТЧЕТ", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, no_indent=True)
    _add_styled_paragraph(
        doc,
        f"о лабораторной работе № {data.get('lab_number', 1)}",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        no_indent=True,
    )
    for _ in range(2):
        doc.add_paragraph()
    if data.get("student_group"):
        _add_styled_paragraph(
            doc,
            f"Выполнил: студент гр. {data['student_group']}",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            no_indent=True,
        )
    if data.get("student_name"):
        _add_styled_paragraph(doc, data["student_name"], alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True)
    if data.get("teacher"):
        _add_styled_paragraph(
            doc,
            f"Проверил: {data['teacher']}",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            no_indent=True,
        )
    for _ in range(2):
        doc.add_paragraph()
    _add_styled_paragraph(doc, data.get("city", "Екатеринбург"), alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True)
    _add_styled_paragraph(doc, data.get("year", "2026"), alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True)


def build_referat(doc: Document, data: dict[str, Any]) -> None:
    """Build the abstract (Реферат) section with summary statistics.

    Args:
        doc: The active Document instance.
        data: The full report dictionary.
    """
    doc.add_page_break()
    _add_styled_paragraph(doc, "РЕФЕРАТ", style="Heading 1")
    tasks = data.get("tasks", [])
    num_tables = sum(1 for t in tasks if t.get("test_data", {}).get("rows"))
    num_listings = sum(
        len(t.get("code_files", [])) or (1 if t.get("code") else 0) for t in tasks
    )
    _add_styled_paragraph(doc, f"Отчет XX с., {num_tables} табл., {num_listings} листингов.")
    if data.get("student_name"):
        _add_styled_paragraph(doc, data["student_name"])
    kw = data.get("keywords", [])
    if kw:
        _add_styled_paragraph(doc, ", ".join(kw).upper())
    for field_key in ("object_of_study", "goal", "methods", "scope"):
        label_map = {
            "object_of_study": "Объект исследования",
            "goal": "Цель работы",
            "methods": "Методы исследования",
            "scope": "Область применения",
        }
        if data.get(field_key):
            _add_styled_paragraph(doc, f"{label_map[field_key]}: {data[field_key]}.")


def build_contents(doc: Document, data: dict[str, Any]) -> None:
    """Insert the table of contents page.

    Args:
        doc: The active Document instance.
        data: Unused here, included for API consistency.
    """
    doc.add_page_break()
    _add_styled_paragraph(doc, "СОДЕРЖАНИЕ", alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True)
    create_toc(doc)


def build_introduction(doc: Document, data: dict[str, Any]) -> None:
    """Build the introduction (Введение) section.

    Args:
        doc: The active Document instance.
        data: The full report dictionary.
    """
    doc.add_page_break()
    _add_styled_paragraph(doc, "ВВЕДЕНИЕ", style="Heading 1")
    intro = data.get("introduction", "Введение...")
    for p in intro.split("\n"):
        if p.strip():
            _add_styled_paragraph(doc, p.strip())


def build_main_part(doc: Document, data: dict[str, Any]) -> None:
    """Build the core body (Основная часть) with per-task subsections.

    Each task is rendered as:
        N.M Заголовок
        N.M.1 Постановка задачи
        N.M.2 Тестовые данные (table)
        N.M.3 Алгоритм и листинг программы

    Args:
        doc: The active Document instance.
        data: The full report dictionary.
    """
    doc.add_page_break()
    _add_styled_paragraph(doc, "ОСНОВНАЯ ЧАСТЬ", style="Heading 1")
    lab_n = data.get("lab_number", 1)

    for task in data.get("tasks", []):
        t_num = task.get("number", 1)
        prefix = f"{lab_n}.{t_num}"

        # --- Task title ---
        _add_styled_paragraph(doc, f"{prefix} Задание {t_num}. {task.get('title', '')}", style="Heading 2")

        # --- Problem statement ---
        _add_styled_paragraph(doc, f"{prefix}.1 Постановка задачи", style="Heading 3")
        for p in task.get("problem_statement", "").split("\n"):
            if p.strip():
                _add_styled_paragraph(doc, p.strip())

        # --- Test data ---
        _add_styled_paragraph(doc, f"{prefix}.2 Тестовые данные", style="Heading 3")
        td = task.get("test_data", {})
        rows = td.get("rows", [])
        if rows:
            caption = td.get("caption", "Данные")
            _add_styled_paragraph(
                doc, f"Таблица {prefix} — {caption}", alignment=WD_ALIGN_PARAGRAPH.CENTER, no_indent=True
            )
            headers = td.get("headers", ["Ввод", "Вывод"])
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
            for r_idx, row in enumerate(rows):
                for c_idx, val in enumerate(row):
                    table.rows[r_idx + 1].cells[c_idx].text = str(val)

        # --- Algorithm & listing ---
        _add_styled_paragraph(doc, f"{prefix}.3 Алгоритм и листинг программы", style="Heading 3")
        for p in task.get("algorithm", "").split("\n"):
            if p.strip():
                _add_styled_paragraph(doc, p.strip())

        code_files = task.get("code_files", [])
        if not code_files and task.get("code"):
            code_files = [{"content": task["code"], "caption": task.get("listing_caption")}]

        for idx, cf in enumerate(code_files):
            for line in cf.get("content", "").split("\n"):
                _add_code_paragraph(doc, line)
            listing_prefix = f"{prefix}.{idx + 1}" if len(code_files) > 1 else prefix
            _add_styled_paragraph(
                doc,
                f"Листинг {listing_prefix} — {cf.get('caption') or 'Программа'}",
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                no_indent=True,
            )

        # --- Optional image ---
        img_path = task.get("image_path")
        if img_path and os.path.exists(img_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = 0
            p.add_run().add_picture(img_path, width=Cm(15))
            _add_styled_paragraph(
                doc,
                f"Рисунок {prefix} — {task.get('image_caption', 'Схема')}",
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                no_indent=True,
            )


def build_conclusion(doc: Document, data: dict[str, Any]) -> None:
    """Build the conclusion (Заключение) section.

    Args:
        doc: The active Document instance.
        data: The full report dictionary.
    """
    doc.add_page_break()
    _add_styled_paragraph(doc, "ЗАКЛЮЧЕНИЕ", style="Heading 1")
    for p in data.get("conclusion", "Задачи выполнены.").split("\n"):
        if p.strip():
            _add_styled_paragraph(doc, p.strip())


def build_extra(doc: Document, data: dict[str, Any]) -> None:
    """Build optional sections: bibliography (Список литературы) and appendices.

    Args:
        doc: The active Document instance.
        data: The full report dictionary.
    """
    bib = data.get("bibliography", [])
    if bib:
        doc.add_page_break()
        _add_styled_paragraph(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", style="Heading 1")
        for i, item in enumerate(bib):
            _add_styled_paragraph(doc, f"{i + 1}. {item}")

    for i, appendix in enumerate(data.get("appendix", [])):
        doc.add_page_break()
        _add_styled_paragraph(
            doc,
            appendix.get("title", f"Приложение {i + 1}"),
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            no_indent=True,
        )
        for p in appendix.get("content", "").split("\n"):
            if p.strip():
                _add_styled_paragraph(doc, p.strip())


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    """Parse command-line arguments.

    Args:
        argv: Argument list (defaults to ``sys.argv[1:]``).

    Returns:
        Parsed namespace with ``input`` and ``output`` paths.
    """
    parser = argparse.ArgumentParser(
        description="Generate a GOST 7.32-2001 compliant .docx lab report from JSON input.",
    )
    parser.add_argument("--input", "-i", required=True, help="Path to input data.json file")
    parser.add_argument("--output", "-o", required=True, help="Path for the generated .docx output")
    return parser.parse_args(argv)


def load_data(path: str) -> dict[str, Any]:
    """Load and validate the JSON report data.

    Args:
        path: Path to the JSON file.

    Returns:
        Parsed dictionary with report data.

    Raises:
        FileNotFoundError: If the file does not exist.
        json.JSONDecodeError: If the file contains invalid JSON.
        ValueError: If required keys are missing.
    """
    path_obj = Path(path)
    if not path_obj.exists():
        raise FileNotFoundError(f"Input file not found: {path}")

    with path_obj.open("r", encoding="utf-8") as f:
        data = json.load(f)

    for key in REQUIRED_JSON_KEYS:
        if key not in data:
            raise ValueError(f"Missing required key '{key}' in {path}")

    return data


def build_report(data: dict[str, Any]) -> Document:
    """Assemble the complete report document from structured data.

    Args:
        data: The full report dictionary.

    Returns:
        A fully built python-docx Document ready to save.
    """
    doc = Document()

    # Remove the default empty paragraph created by python-docx
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    setup_gost_styles(doc)
    configure_section(doc)

    build_title_page(doc, data)
    build_referat(doc, data)
    build_contents(doc, data)
    build_introduction(doc, data)
    build_main_part(doc, data)
    build_conclusion(doc, data)
    build_extra(doc, data)

    return doc


def main(argv: list[str] | None = None) -> int:
    """CLI entry point.

    Args:
        argv: Command-line arguments (defaults to ``sys.argv[1:]``).

    Returns:
        Exit code (0 on success, 1 on error).
    """
    args = parse_args(argv)

    try:
        data = load_data(args.input)
    except (FileNotFoundError, json.JSONDecodeError, ValueError) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1

    doc = build_report(data)

    try:
        doc.save(args.output)
    except OSError as exc:
        print(f"ERROR: Failed to save document: {exc}", file=sys.stderr)
        return 1

    print(f"Report generated: {args.output}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
